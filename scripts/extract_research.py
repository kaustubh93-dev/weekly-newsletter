#!/usr/bin/env python3
"""
extract_research.py — Extract text from all files in a research topic folder.

Reads PDFs, DOCX, TXT, MD, and link files from research/{topic}/ and produces
a single consolidated markdown file at output/{topic}/_extracted.md.

Usage:
    python scripts/extract_research.py --topic mastering-windows-event-logs
    python scripts/extract_research.py --topic pagefile-deep-dive --list

Prerequisites:
    pip install -r scripts/requirements.txt
"""

import argparse
import os
import re
import sys
from datetime import datetime


def extract_pdf(filepath: str) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(filepath)
        text_parts = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(f"### Page {i + 1}\n{page_text}")
        return "\n\n".join(text_parts) if text_parts else "(No extractable text found in PDF)"
    except ImportError:
        return "(⚠️ PyPDF2 not installed — run: pip install -r scripts/requirements.txt)"
    except Exception as e:
        return f"(⚠️ Error reading PDF: {e})"


def extract_docx(filepath: str) -> str:
    """Extract text from a Word document."""
    try:
        from docx import Document
        doc = Document(filepath)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract tables
        table_texts = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                header = rows[0]
                separator = " | ".join(["---"] * len(table.rows[0].cells))
                table_text = "\n".join([header, separator] + rows[1:])
                table_texts.append(table_text)

        result = "\n\n".join(paragraphs)
        if table_texts:
            result += "\n\n#### Tables Found:\n\n" + "\n\n".join(table_texts)
        return result if result.strip() else "(No extractable text found in DOCX)"
    except ImportError:
        return "(⚠️ python-docx not installed — run: pip install -r scripts/requirements.txt)"
    except Exception as e:
        return f"(⚠️ Error reading DOCX: {e})"


def extract_text_file(filepath: str) -> str:
    """Read a plain text or markdown file."""
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    except Exception as e:
        return f"(⚠️ Error reading file: {e})"


def extract_image_info(filepath: str) -> str:
    """Return image metadata (can't extract text without OCR)."""
    size_kb = os.path.getsize(filepath) / 1024
    return f"📸 Image file: {os.path.basename(filepath)} ({size_kb:.0f} KB)\n(Image content — describe manually or use OCR)"


def get_file_handler(filepath: str):
    """Return the appropriate extraction function based on file extension."""
    ext = os.path.splitext(filepath)[1].lower()
    handlers = {
        ".pdf": extract_pdf,
        ".docx": extract_docx,
        ".doc": extract_docx,
        ".txt": extract_text_file,
        ".md": extract_text_file,
        ".markdown": extract_text_file,
        ".url": extract_text_file,
        ".link": extract_text_file,
        ".png": extract_image_info,
        ".jpg": extract_image_info,
        ".jpeg": extract_image_info,
        ".gif": extract_image_info,
        ".bmp": extract_image_info,
        ".webp": extract_image_info,
    }
    return handlers.get(ext)


def extract_topic(topic: str, repo_root: str) -> str:
    """Extract text from all files in a research topic folder."""
    research_path = os.path.join(repo_root, "research", topic)

    if not os.path.exists(research_path):
        print(f"❌ Research folder not found: {research_path}")
        print(f"   Create it with: mkdir research\\{topic}")
        sys.exit(1)

    # Collect all files
    files = []
    for root, dirs, filenames in os.walk(research_path):
        for filename in sorted(filenames):
            filepath = os.path.join(root, filename)
            if not filename.startswith("."):
                files.append(filepath)

    if not files:
        print(f"❌ No files found in: {research_path}")
        print(f"   Drop your PDFs, docs, screenshots, and notes there first.")
        sys.exit(1)

    # Extract content from each file
    sections = []
    sections.append(f"# Research Extraction: {topic}")
    sections.append(f"> Extracted on {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    sections.append(f"> Source folder: research/{topic}/")
    sections.append(f"> Total files: {len(files)}")
    sections.append("---\n")

    extracted_count = 0
    skipped_count = 0

    for filepath in files:
        filename = os.path.basename(filepath)
        rel_path = os.path.relpath(filepath, research_path)
        handler = get_file_handler(filepath)

        if handler:
            print(f"  📄 Extracting: {rel_path}")
            content = handler(filepath)
            sections.append(f"## 📄 {rel_path}\n")
            sections.append(content)
            sections.append("\n---\n")
            extracted_count += 1
        else:
            print(f"  ⏭️  Skipping unsupported: {rel_path}")
            sections.append(f"## ⏭️ {rel_path}\n")
            sections.append(f"(Unsupported file type: {os.path.splitext(filename)[1]})\n")
            sections.append("---\n")
            skipped_count += 1

    # Add summary at the end
    sections.append(f"\n## 📊 Extraction Summary")
    sections.append(f"- Files processed: {extracted_count}")
    sections.append(f"- Files skipped: {skipped_count}")
    sections.append(f"- Total files: {len(files)}")

    return "\n\n".join(sections), extracted_count, skipped_count


def list_topics(repo_root: str):
    """List all research topics."""
    research_path = os.path.join(repo_root, "research")
    if not os.path.exists(research_path):
        print("No research/ folder found.")
        return

    topics = [d for d in os.listdir(research_path)
              if os.path.isdir(os.path.join(research_path, d))]

    if not topics:
        print("No research topics found. Create one with:")
        print("  mkdir research\\your-topic-name")
        return

    print(f"📂 Research topics ({len(topics)}):\n")
    for topic in sorted(topics):
        topic_path = os.path.join(research_path, topic)
        files = [f for f in os.listdir(topic_path) if not f.startswith(".")]
        output_exists = os.path.exists(os.path.join(repo_root, "output", topic, "_extracted.md"))
        status = "✅ extracted" if output_exists else "⏳ not yet extracted"
        print(f"  📁 {topic}/ ({len(files)} files) — {status}")


def main():
    parser = argparse.ArgumentParser(
        description="Extract text from research materials into a consolidated markdown file",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python scripts/extract_research.py --topic mastering-windows-event-logs
  python scripts/extract_research.py --list
        """
    )
    parser.add_argument("--topic", help="Topic folder name inside research/")
    parser.add_argument("--list", action="store_true", help="List all research topics")

    args = parser.parse_args()
    repo_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))

    if args.list:
        list_topics(repo_root)
        return

    if not args.topic:
        parser.error("--topic is required (or use --list)")

    topic = args.topic.strip("/\\")

    print(f"📚 Extracting research: {topic}")
    print(f"{'─' * 50}")

    content, extracted, skipped = extract_topic(topic, repo_root)

    # Save to output folder
    output_dir = os.path.join(repo_root, "output", topic)
    os.makedirs(output_dir, exist_ok=True)
    output_file = os.path.join(output_dir, "_extracted.md")

    with open(output_file, "w", encoding="utf-8") as f:
        f.write(content)

    print(f"\n{'─' * 50}")
    print(f"✅ Extraction complete!")
    print(f"   📄 Extracted: {extracted} files")
    print(f"   ⏭️  Skipped: {skipped} files")
    print(f"   📝 Output: output/{topic}/_extracted.md")
    print(f"\n📋 Next steps:")
    print(f"   1. Open output/{topic}/_extracted.md")
    print(f"   2. Feed it to Copilot CLI along with the prompts in prompts/")
    print(f"   3. Save the 3 outputs into output/{topic}/:")
    print(f"      • newsletter.html")
    print(f"      • blog-post.md")
    print(f"      • linkedin-post.txt")
    print(f"   4. Run: python scripts/prepare_edition.py --topic {topic} --week YYYY-WNN")


if __name__ == "__main__":
    main()
